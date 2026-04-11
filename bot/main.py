import os
import re
import time as _time_module
import asyncio
import random
import threading
import functools

from telethon import TelegramClient, events, Button
from telethon.errors import FloodWaitError, AlreadyInConversationError
from telethon.tl.functions.updates import GetStateRequest

from config import (
    API_ID,
    API_HASH,
    BOT_TOKEN,
    OWNER_ID,
    SESSIONS_DIR,
    CHANNEL_KEYS,
)
from database import (
    load_db,
    save_db,
    clear_seen,
    get_seen_count,
    get_raw_count,
    load_raw_links,
    save_raw_links,
    normalize_link,
    load_all_known_links,
)
from account_manager import AccountManager
from channel_setup import (
    create_archive_channels,
    add_account_to_channels,
    add_owner_to_channels,
    join_accounts_via_invites,
    _first_authorized_session,
)
from telethon.tl.functions.channels import GetParticipantRequest
from telethon.errors import UserNotParticipantError, ChannelPrivateError
from harvester import harvest_sources
from sorter import run_sorter, clear_archive_channels, sort_links_inline
from joiner import run_smart_joiner
from searcher import run_smart_discovery
import state as sorter_ctrl

db = load_db()
bot = TelegramClient(
    "bot_controller", API_ID, API_HASH,
    connection_retries=-1,
    retry_delay=5,
    auto_reconnect=True,
    request_retries=5,
)


# ─────────────────────────────────────────────────────────────────────────────
# Auth guard
# ─────────────────────────────────────────────────────────────────────────────

def _is_authorized(sender_id: int) -> bool:
    """True if sender is the owner or a trusted user."""
    fresh = load_db()
    return sender_id == OWNER_ID or sender_id in fresh.get("trusted_users", [])


# Timestamp set in main() the moment the bot successfully starts.
# Any event older than this is considered a stale replay and is silently dropped.
_BOT_START_TIME: float = 0.0


def _is_stale_event(event) -> bool:
    """
    True if the event pre-dates the current bot session.
    Covers both NewMessage and CallbackQuery events so stale replays
    after a restart are silently discarded instead of flooding the chat.
    """
    if not _BOT_START_TIME:
        return False
    try:
        if isinstance(event, events.CallbackQuery.Event):
            ts_raw = getattr(event.query, "date", None) or getattr(event, "date", None)
            ts = float(ts_raw) if ts_raw else None
        elif isinstance(event, events.NewMessage.Event):
            date = getattr(getattr(event, "message", None), "date", None)
            if date is None:
                return False
            ts = date.timestamp() if hasattr(date, "timestamp") else float(date)
        else:
            return False
        if ts and ts < _BOT_START_TIME:
            return True
    except Exception:
        pass
    return False


# Keep old name as alias so callers that weren't updated still work
_is_stale_callback = _is_stale_event


def owner_only(func):
    """Allows both the owner AND trusted users."""
    @functools.wraps(func)
    async def wrapper(event):
        # Silently drop stale callback queries — prevents message flooding on restart
        if _is_stale_callback(event):
            try:
                await event.answer()
            except Exception:
                pass
            return

        if not _is_authorized(event.sender_id):
            msg = "🚫 غير مصرح لك. أرسل /start لطلب الوصول."
            is_callback = isinstance(event, events.CallbackQuery.Event)
            if is_callback:
                # Only answer (ephemeral alert) — never send a persistent message
                try:
                    await event.answer(msg, alert=True)
                except Exception:
                    pass
            else:
                # Message event: a single respond is fine
                try:
                    await event.respond(msg, parse_mode="md")
                except Exception:
                    pass
            return
        try:
            await func(event)
        except AlreadyInConversationError:
            try:
                await event.answer("⚠️ هناك عملية جارية. أنهها أو انتظر.", alert=True)
            except Exception:
                pass
        except FloodWaitError as e:
            print(f"[FloodWait] {func.__name__}: {e.seconds}s — skipping response")
        except Exception as e:
            try:
                await event.answer("❌ حدث خطأ، حاول مجدداً.", alert=True)
            except Exception:
                pass
            try:
                await bot.send_message(OWNER_ID, f"❌ خطأ في {func.__name__}:\n{e}", parse_mode="md")
            except Exception:
                pass
    return wrapper


def admin_only(func):
    """Strict owner-only — trusted users cannot access."""
    @functools.wraps(func)
    async def wrapper(event):
        # Silently drop stale callback queries — prevents message flooding on restart
        if _is_stale_callback(event):
            try:
                await event.answer()
            except Exception:
                pass
            return

        if event.sender_id != OWNER_ID:
            msg = "🔒 هذا الإجراء للمالك فقط."
            is_callback = isinstance(event, events.CallbackQuery.Event)
            if is_callback:
                try:
                    await event.answer(msg, alert=True)
                except Exception:
                    pass
            else:
                try:
                    await event.respond(msg, parse_mode="md")
                except Exception:
                    pass
            return
        try:
            await func(event)
        except AlreadyInConversationError:
            try:
                await event.answer("⚠️ هناك عملية جارية. أنهها أو انتظر.", alert=True)
            except Exception:
                pass
        except FloodWaitError as e:
            print(f"[FloodWait] {func.__name__}: {e.seconds}s — skipping response")
        except Exception as e:
            try:
                await event.answer("❌ حدث خطأ، حاول مجدداً.", alert=True)
            except Exception:
                pass
            try:
                await bot.send_message(OWNER_ID, f"❌ خطأ في {func.__name__}:\n{e}", parse_mode="md")
            except Exception:
                pass
    return wrapper


async def status_msg(text: str):
    try:
        await bot.send_message(OWNER_ID, text, parse_mode="md")
    except Exception:
        pass


async def _get_authorized_sessions(accounts: list) -> list:
    """Return only sessions that are currently authorized (connected accounts)."""
    authorized = []
    for sess in accounts:
        from telethon import TelegramClient as _TC
        client = _TC(sess, API_ID, API_HASH)
        try:
            await client.connect()
            if await client.is_user_authorized():
                authorized.append(sess)
        except Exception:
            pass
        finally:
            try:
                await client.disconnect()
            except Exception:
                pass
    return authorized


async def _check_and_notify_membership(session: str, db: dict):
    """
    For a newly added account: check membership in all archive channels and
    source groups. Send the account a message listing channels/groups it needs
    to join (with their links/IDs).
    """
    channels = db.get("channels", {})
    sources  = db.get("sources", [])
    channel_invites = db.get("channels_invites", [])

    if not channels and not sources:
        return

    from telethon import TelegramClient as _TC
    from telethon.tl.types import Channel as _TLChannel, Chat as _TLChat
    client = _TC(session, API_ID, API_HASH)
    try:
        await client.connect()
        if not await client.is_user_authorized():
            return

        missing_channels = []
        missing_sources  = []

        # Check archive channels
        for key, ch_id in channels.items():
            if not isinstance(ch_id, int):
                continue
            from config import CHANNEL_KEYS as _CK
            label = _CK.get(key, key)
            is_member = False
            try:
                me = await client.get_me()
                await client(GetParticipantRequest(ch_id, me.id))
                is_member = True
            except (UserNotParticipantError, ChannelPrivateError):
                is_member = False
            except Exception:
                is_member = False

            if not is_member:
                # Try to find invite link
                invite = None
                for lnk in channel_invites:
                    missing_channels.append((label, lnk))
                    invite = lnk
                    break
                if not invite:
                    missing_channels.append((label, f"ID: `{ch_id}`"))

        # Check source groups
        for src in sources:
            is_member = False
            try:
                entity = await client.get_entity(src)
                me = await client.get_me()
                if isinstance(entity, (_TLChannel, _TLChat)):
                    await client(GetParticipantRequest(entity, me.id))
                    is_member = True
            except (UserNotParticipantError,):
                is_member = False
            except Exception:
                is_member = False

            if not is_member:
                missing_sources.append(src)

        if missing_channels or missing_sources:
            lines = ["📋 **مرحباً! الحساب تمت إضافته للبوت.**\n\nيجب الانضمام للقنوات والمجموعات التالية:"]

            if missing_channels:
                lines.append("\n🗂 **قنوات الأرشيف (يجب الانضمام):**")
                seen_links = set()
                for label, link in missing_channels:
                    if link not in seen_links:
                        lines.append(f"• {label}: {link}")
                        seen_links.add(link)

            if missing_sources:
                lines.append("\n🔗 **مجموعات المصادر (يجب الانضمام):**")
                for src in missing_sources:
                    lines.append(f"• {src}")

            try:
                me = await client.get_me()
                await client.send_message(me.id, "\n".join(lines), parse_mode="md")
            except Exception:
                pass

    except Exception:
        pass
    finally:
        try:
            await client.disconnect()
        except Exception:
            pass


def make_edit_callback(msg_id: int, chat_id: int, fixed_buttons=None):
    """Returns an async callback that EDITS one persistent message instead of sending new ones."""
    async def _cb(text: str):
        try:
            await bot.edit_message(
                chat_id, msg_id, text,
                buttons=fixed_buttons,
                parse_mode="md",
            )
        except Exception:
            pass
    return _cb


async def _keep_alive_http():
    """
    Minimal HTTP health-check server — keeps autoscale deployment alive.

    Replit sets REPLIT_DEPLOYMENT_ID only in the production deployment
    environment, not in the dev workspace. So we skip the server entirely
    in development — the dev workspace has the API server on the same port
    and the keep_alive.sh script already handles bot restarts.
    """
    is_replit_deploy = bool(os.environ.get("REPLIT_DEPLOYMENT_ID"))
    is_railway       = bool(os.environ.get("RAILWAY_ENVIRONMENT"))
    if not is_replit_deploy and not is_railway:
        print("ℹ️ Dev workspace detected — HTTP health-check server not started.")
        return

    port = int(os.environ.get("PORT", 8080))

    async def _handle(reader, writer):
        try:
            await reader.read(1024)
        except Exception:
            pass
        body = b"Bot OK"
        response = (
            b"HTTP/1.1 200 OK\r\nContent-Type: text/plain\r\nContent-Length: "
            + str(len(body)).encode()
            + b"\r\n\r\n" + body
        )
        try:
            writer.write(response)
            await writer.drain()
        except Exception:
            pass
        writer.close()

    try:
        server = await asyncio.start_server(_handle, "0.0.0.0", port)
        print(f"🌐 HTTP health-check server started on port {port}")
        async with server:
            await server.serve_forever()
    except Exception as e:
        print(f"⚠️ HTTP server error: {e}")


# ─────────────────────────────────────────────────────────────────────────────
# Smart Flow Engine — determines current progress & next recommended step
# ─────────────────────────────────────────────────────────────────────────────

def get_flow_status(db: dict) -> dict:
    accounts   = db.get("accounts", [])
    channels   = db.get("channels", {})
    sources    = db.get("sources", [])
    raw_count  = get_raw_count()
    seen_count = get_seen_count()
    stats      = db.get("stats", {})
    sorted_ok  = stats.get("total_sorted", 0)
    last_idx   = db.get("progress", {}).get("last_sorted_index", 0)

    steps = [
        {
            "n": 1, "key": "accounts",
            "icon": "✅" if accounts else "🔴",
            "label": "الحسابات",
            "detail": f"{len(accounts)} حساب مرتبط" if accounts else "لا يوجد حساب",
            "done": bool(accounts),
        },
        {
            "n": 2, "key": "channels",
            "icon": "✅" if len(channels) >= 7 else ("⚠️" if channels else "🔴"),
            "label": "قنوات الأرشيف",
            "detail": f"{len(channels)}/7 قنوات" if channels else "لم تُنشأ بعد",
            "done": len(channels) >= 7,
            "locked": not accounts,
        },
        {
            "n": 3, "key": "sources",
            "icon": "✅" if sources else "🔴",
            "label": "المصادر",
            "detail": f"{len(sources)} مصدر مضاف" if sources else "لا مصادر",
            "done": bool(sources),
            "locked": not accounts,
        },
        {
            "n": 4, "key": "harvest",
            "icon": "✅" if raw_count > 0 else "🔴",
            "label": "حصاد الروابط",
            "detail": f"{raw_count:,} رابط تم جمعها" if raw_count else "لم يبدأ بعد",
            "done": raw_count > 0,
            "locked": not sources,
        },
        {
            "n": 5, "key": "sort",
            "icon": "✅" if (raw_count > 0 and last_idx >= raw_count) else ("🔄" if last_idx > 0 else "⏳"),
            "label": "الفرز الشامل",
            "detail": (
                f"{last_idx:,}/{raw_count:,} ✦ {int(last_idx/raw_count*100)}%"
                if raw_count > 0 and last_idx > 0
                else ("اكتمل ✅" if raw_count > 0 and sorted_ok > 0 and last_idx >= raw_count else "لم يبدأ بعد")
            ),
            "done": raw_count > 0 and last_idx >= raw_count,
            "in_progress": raw_count > 0 and 0 < last_idx < raw_count,
            "locked": raw_count == 0,
        },
        {
            "n": 6, "key": "discover",
            "icon": "✅" if stats.get("total_found", 0) > raw_count else "⏳",
            "label": "اكتشاف ذكي",
            "detail": "يجد روابط طبية تلقائياً" if not stats.get("total_found") else "تم التشغيل",
            "done": False,
            "locked": not accounts,
        },
        {
            "n": 7, "key": "join",
            "icon": "✅" if db.get("joined_links") else "⏳",
            "label": "انضمام ذكي",
            "detail": f"{len(db.get('joined_links', []))} رابط تم الانضمام" if db.get("joined_links") else "لم يبدأ بعد",
            "done": bool(db.get("joined_links")),
            "locked": not channels,
        },
    ]

    # Determine recommended next step
    next_step = None
    for s in steps:
        if not s.get("locked") and not s["done"]:
            next_step = s
            break
        if s.get("in_progress"):
            next_step = s
            break

    return {"steps": steps, "next": next_step, "last_idx": last_idx, "raw_count": raw_count}


def build_dashboard(db: dict) -> tuple[str, list]:
    flow   = get_flow_status(db)
    steps  = flow["steps"]
    nxt    = flow["next"]
    last   = flow["last_idx"]

    done_count = sum(1 for s in steps if s["done"])
    pct        = int(done_count / 7 * 100)
    bar        = "█" * done_count + "░" * (7 - done_count)

    # ── Compact step list ─────────────────────────────────────────────────────
    step_lines = []
    for s in steps:
        lock = "🔒" if s.get("locked") else ""
        step_lines.append(f"{s['icon']} {s['n']}. {lock}{s['label']} — {s['detail']}")

    text = (
        "🏥 **نظام الفلترة الطبية الذكي**\n"
        f"📊 {bar} {pct}%\n\n"
        + "\n".join(step_lines)
    )

    if nxt:
        tip = {
            "accounts": "ربط حساب تيليجرام",
            "channels": "إنشاء قنوات الأرشيف",
            "sources":  "إضافة مجموعات كمصادر",
            "harvest":  "جمع الروابط من المصادر",
            "sort":     f"استئناف من {last+1:,}" if last > 0 else "فرز وتصنيف الروابط",
            "discover": "اكتشاف روابط طبية جديدة",
            "join":     "الانضمام للقنوات الطبية",
        }.get(nxt["key"], "")
        text += f"\n\n💡 **التالي:** {nxt['label']} — {tip}"
    else:
        text += "\n\n🎉 **جميع الخطوات مكتملة!**"

    # ── Buttons ───────────────────────────────────────────────────────────────
    rows = []

    # Highlighted next step
    next_btn_map = {
        "accounts": ("➕ ربط حساب ◄",    b"add_acc"),
        "channels": ("📺 إنشاء القنوات ◄", b"make_ch"),
        "sources":  ("🔗 إضافة مصدر ◄",   b"add_src"),
        "harvest":  ("🌾 بدء الحصاد ◄",    b"harvest"),
        "sort":     ("▶️ استئناف الفرز ◄" if last > 0 else "⚡ بدء الفرز ◄", b"run_sort"),
        "discover": ("🧠 اكتشاف ذكي ◄",    b"smart_discover"),
        "join":     ("🤝 انضمام ذكي ◄",    b"smart_join"),
    }
    if nxt and nxt["key"] in next_btn_map:
        lbl, dat = next_btn_map[nxt["key"]]
        rows.append([Button.inline(lbl, dat)])

    rows += [
        [Button.inline("➕ حساب", b"add_acc"),   Button.inline("👤 حساباتي", b"list_acc"),  Button.inline("📊 إحصائيات", b"stats")],
        [Button.inline("📺 قنوات", b"make_ch"),  Button.inline("🔗 مصادر",   b"list_src"),  Button.inline("✏️ أضف مصدر", b"add_src")],
        [Button.inline("🌾 حصاد", b"harvest"),   Button.inline("⚡ فرز",     b"run_sort"),  Button.inline("🧠 اكتشاف",  b"smart_discover")],
        [Button.inline("🤝 انضمام", b"smart_join"), Button.inline("🧹 مسح الذاكرة", b"clear_mem")],
    ]

    return text, rows


# ─────────────────────────────────────────────────────────────────────────────
# Nav helpers — reusable nav row + post-step guidance
# ─────────────────────────────────────────────────────────────────────────────

def nav_row(prev_step_btn: bytes | None = None) -> list:
    """Bottom navigation: [◀ back step] [🏠 home]."""
    row = []
    if prev_step_btn:
        row.append(Button.inline("◀️ رجوع", prev_step_btn))
    row.append(Button.inline("🏠 القائمة الرئيسية", b"home"))
    return row


async def send_next_step_hint(next_key: str, db: dict):
    """After completing a step, tell user what to do next."""
    flow = get_flow_status(db)
    nxt  = flow["next"]
    if not nxt:
        await status_msg(
            "🎉 **جميع الخطوات مكتملة!**\n"
            "البوت جاهز بالكامل. استخدم /start للوحة التحكم."
        )
        return

    hints = {
        "channels":  ("📺 القنوات",      b"make_ch"),
        "sources":   ("🔗 إضافة مصدر",   b"add_src"),
        "harvest":   ("🌾 بدء الحصاد",   b"harvest"),
        "sort":      ("⚡ بدء الفرز",    b"run_sort"),
        "discover":  ("🧠 اكتشاف ذكي",   b"smart_discover"),
        "join":      ("🤝 انضمام ذكي",    b"smart_join"),
    }
    if nxt["key"] in hints:
        lbl, dat = hints[nxt["key"]]
        await bot.send_message(
            OWNER_ID,
            f"✅ **تم!** الخطوة التالية الموصى بها:\n\n"
            f"**{nxt['n']}. {nxt['label']}** — _{nxt['detail']}_",
            buttons=[[Button.inline(f"⏭️ {lbl}", dat)], nav_row()],
            parse_mode="md",
        )


# ─────────────────────────────────────────────────────────────────────────────
# /start & 🏠 Home
# ─────────────────────────────────────────────────────────────────────────────

async def show_dashboard(target):
    """Send or edit the dashboard. target can be a message event or a callback event."""
    db.update(load_db())
    text, buttons = build_dashboard(db)
    # Try to edit existing message first (cleaner UX), fall back to new message
    try:
        await target.edit(text, buttons=buttons, parse_mode="md")
    except Exception:
        try:
            await target.respond(text, buttons=buttons, parse_mode="md")
        except Exception:
            await bot.send_message(OWNER_ID, text, buttons=buttons, parse_mode="md")


@bot.on(events.NewMessage(pattern="/start"))
async def start_handler(event):
    if _is_stale_event(event):
        return
    db.update(load_db())
    try:
        if _is_authorized(event.sender_id):
            text, buttons = build_dashboard(db)
            await event.respond(text, buttons=buttons, parse_mode="md")
            return

        user_id_str = str(event.sender_id)
        if user_id_str in db.get("pending_requests", {}):
            await event.respond(
                "⏳ **طلبك قيد المراجعة**\n\nسيتم إشعارك فور قبول أو رفض طلبك من المالك.",
                parse_mode="md",
            )
            return

        sender = await event.get_sender()
        name = getattr(sender, "first_name", "") or ""
        await event.respond(
            f"👋 مرحباً {name}!\n\n"
            "هذا البوت خاص ويتطلب موافقة المالك للوصول.\n\n"
            "هل تريد إرسال طلب وصول للمالك؟",
            buttons=[[Button.inline("📨 إرسال طلب الوصول", f"req_{event.sender_id}".encode())]],
            parse_mode="md",
        )
    except FloodWaitError as e:
        print(f"[FloodWait] start_handler: {e.seconds}s — skipping response")


@bot.on(events.CallbackQuery(data=b"home"))
@owner_only
async def home_handler(event):
    await event.answer()
    await show_dashboard(event)


# ─────────────────────────────────────────────────────────────────────────────
# Step 1 — Add account
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"add_acc"))
@admin_only
async def add_acc_handler(event):
    await event.answer()
    count = len(db.get("accounts", []))

    # ── Single persistent message — all updates go through event.edit() ──────
    await event.edit(
        f"**①  ربط حساب تيليجرام**\n"
        f"━━━━━━━━━━━━━━━━━━━━━\n"
        f"الحسابات المرتبطة حالياً: **{count}**\n\n"
        f"📱 أرسل رقم الهاتف بالصيغة الدولية:\n"
        f"`+9671234567890`",
        buttons=[nav_row()],
        parse_mode="md",
    )

    msg_id  = event.message_id
    chat_id = OWNER_ID

    async def _edit(text: str, buttons=None):
        try:
            await bot.edit_message(chat_id, msg_id, text, buttons=buttons or [nav_row()], parse_mode="md")
        except Exception:
            pass

    try:
        async with bot.conversation(OWNER_ID, timeout=120, exclusive=False) as conv:
            async def _get_input():
                ev = await conv.wait_event(events.NewMessage(incoming=True, from_users=OWNER_ID))
                return ev.message

            phone_msg = await _get_input()
            phone     = phone_msg.text.strip()
            try:
                await phone_msg.delete()
            except Exception:
                pass

            await _edit("⏳ جاري إرسال كود التحقق...")
            success, result = await AccountManager.add_account_interactive(_get_input, phone, edit_fn=_edit)

            if success:
                is_new = result not in db["accounts"]
                if is_new:
                    db["accounts"].append(result)
                    save_db(db)

                info = await AccountManager.get_account_info(result)
                await _edit(
                    f"✅ **تم ربط الحساب بنجاح!**\n\n"
                    f"👤 {info['name']}  {info['username']}\n"
                    f"☎️ {info['phone']}\n\n"
                    f"💡 لإضافة الحساب لقنوات الأرشيف استخدم **🔗 انضمام عبر روابط دعوة** من قائمة القنوات.",
                    buttons=[[Button.inline("➕ ربط حساب آخر", b"add_acc")], nav_row()],
                )
            else:
                await _edit(
                    f"❌ **فشل ربط الحساب:**\n{result}",
                    buttons=[[Button.inline("🔄 حاول مرة أخرى", b"add_acc")], nav_row()],
                )
    except asyncio.TimeoutError:
        await _edit("⏰ انتهت مهلة الإدخال. اضغط **ربط حساب** للمحاولة مجدداً.",
                    buttons=[[Button.inline("🔄 حاول مرة أخرى", b"add_acc")], nav_row()])
    except AlreadyInConversationError:
        await _edit(
            "⚠️ **يوجد إدخال آخر مفتوح.**\n\n"
            "أرسل أي رسالة لإلغائه، ثم اضغط **ربط حساب** مرة أخرى.",
            buttons=[[Button.inline("🔄 حاول مرة أخرى", b"add_acc")], nav_row()],
        )
    except Exception as e:
        await _edit(
            f"❌ **خطأ غير متوقع:**\n`{e}`\n\nاضغط **حاول مرة أخرى**.",
            buttons=[[Button.inline("🔄 حاول مرة أخرى", b"add_acc")], nav_row()],
        )


# ─────────────────────────────────────────────────────────────────────────────
# List accounts
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"list_acc"))
@owner_only
async def list_acc_handler(event):
    await event.answer()
    if not db["accounts"]:
        await event.edit(
            "❌ لا توجد حسابات مرتبطة بعد.",
            buttons=[[Button.inline("➕ ربط حساب الآن", b"add_acc")], nav_row()],
        )
        return

    lines = [f"👤 **الحسابات المرتبطة ({len(db['accounts'])}):**\n"]
    unauthorized_count = 0
    for i, acc in enumerate(db["accounts"], 1):
        info = await AccountManager.get_account_info(acc)
        if info.get("unauthorized"):
            unauthorized_count += 1
            lines.append(
                f"{i}. **{info['name']}**\n"
                f"   ☎️ {info['phone']}\n"
                f"   ⚠️ _(يجب إعادة ربط هذا الحساب)_"
            )
        else:
            lines.append(f"{i}. **{info['name']}** {info['username']}\n   ☎️ {info['phone']}")

    if unauthorized_count:
        lines.append(
            f"\n🔴 **تحذير:** {unauthorized_count} حساب منتهي الجلسة!\n"
            "سبب شائع: انتقال البوت لخادم جديد (IP مختلف) فأوقف تيليجرام الجلسات تلقائياً.\n"
            "**الحل:** أعد ربط كل حساب منتهٍ عبر زر ➕ ربط حساب."
        )

    # Show which account is the current poster
    poster_session = db.get("poster_session")
    poster_label = "غير محدد ⚠️"
    if poster_session:
        for acc in db["accounts"]:
            if acc == poster_session:
                pi = await AccountManager.get_account_info(acc)
                poster_label = f"{pi['name']} {pi['phone']}"
                break

    lines.append(f"\n📤 **حساب الناشر (أدمن الأرشيف):** {poster_label}")

    buttons = []
    if unauthorized_count:
        buttons.append([Button.inline("➕ إعادة ربط حساب", b"add_acc")])
    else:
        buttons.append([Button.inline("➕ إضافة حساب آخر", b"add_acc")])
    buttons.append([Button.inline("📤 تعيين حساب الناشر", b"set_poster_menu"),
                    Button.inline("🗑 حذف حساب", b"del_acc_menu")])
    buttons.append(nav_row())

    await event.edit(
        "\n".join(lines),
        buttons=buttons,
        parse_mode="md",
    )


# ─────────────────────────────────────────────────────────────────────────────
# Delete account — menu + confirm + execute
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"del_acc_menu"))
@owner_only
async def del_acc_menu_handler(event):
    await event.answer()
    accounts = db.get("accounts", [])
    if not accounts:
        await event.edit(
            "❌ لا توجد حسابات لحذفها.",
            buttons=[nav_row()],
            parse_mode="md",
        )
        return

    lines = ["🗑 **اختر الحساب الذي تريد حذفه:**\n"]
    buttons = []
    for i, acc in enumerate(accounts):
        info = await AccountManager.get_account_info(acc)
        label = f"{info['name']} — {info['phone']}"
        if info.get("unauthorized"):
            label += " ⚠️"
        lines.append(f"{i + 1}. {label}")
        buttons.append([Button.inline(f"🗑 {label}", f"del_acc_confirm_{i}".encode())])

    buttons.append([Button.inline("↩️ رجوع", b"list_acc")])

    await event.edit(
        "\n".join(lines),
        buttons=buttons,
        parse_mode="md",
    )


@bot.on(events.CallbackQuery(data=re.compile(b"del_acc_confirm_(\\d+)")))
@owner_only
async def del_acc_confirm_handler(event):
    await event.answer()
    idx = int(event.data.decode().split("del_acc_confirm_")[1])
    accounts = db.get("accounts", [])
    if idx >= len(accounts):
        await event.edit("⚠️ الحساب غير موجود.", buttons=[nav_row()])
        return

    info = await AccountManager.get_account_info(accounts[idx])
    label = f"{info['name']} — {info['phone']}"

    await event.edit(
        f"⚠️ **تأكيد الحذف**\n\n"
        f"هل أنت متأكد من حذف الحساب:\n**{label}**\n\n"
        f"سيتم إزالته من قائمة الحسابات.\n"
        f"_(لن يتم حذف الجلسة من التيليجرام، فقط من البوت)_",
        buttons=[
            [Button.inline("✅ نعم، احذفه", f"del_acc_do_{idx}".encode()),
             Button.inline("❌ إلغاء", b"del_acc_menu")],
        ],
        parse_mode="md",
    )


@bot.on(events.CallbackQuery(data=re.compile(b"del_acc_do_(\\d+)")))
@owner_only
async def del_acc_do_handler(event):
    await event.answer()
    idx = int(event.data.decode().split("del_acc_do_")[1])
    accounts = db.get("accounts", [])
    if idx >= len(accounts):
        await event.edit("⚠️ الحساب غير موجود.", buttons=[nav_row()])
        return

    session_path = accounts[idx]
    info = await AccountManager.get_account_info(session_path)
    label = f"{info['name']} — {info['phone']}"

    # Remove from db
    db["accounts"].pop(idx)
    save_db(db)

    # Try to delete the session file (non-fatal if it fails)
    # Telethon stores sessions as <path>.session (SQLite) and <path>.session-journal
    try:
        for suffix in (".session", ".session-journal"):
            p = session_path + suffix
            if os.path.exists(p):
                os.remove(p)
    except Exception:
        pass

    remaining = len(db["accounts"])
    await event.edit(
        f"✅ **تم حذف الحساب:**\n{label}\n\n"
        f"الحسابات المتبقية: **{remaining}**",
        buttons=[
            [Button.inline("👤 قائمة الحسابات", b"list_acc")],
            nav_row(),
        ],
        parse_mode="md",
    )


# ─────────────────────────────────────────────────────────────────────────────
# Set poster account (the one with admin rights to archive channels)
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"set_poster_menu"))
@owner_only
async def set_poster_menu_handler(event):
    await event.answer()
    accounts = db.get("accounts", [])
    if not accounts:
        await event.edit("❌ لا توجد حسابات.", buttons=[nav_row()])
        return

    current_poster = db.get("poster_session", "")
    lines = ["📤 **اختر الحساب الناشر (الأدمن في قنوات الأرشيف):**\n",
             "_هذا الحساب هو الوحيد الذي يرسل الروابط إلى قنوات الأرشيف — يجب أن يكون أدمناً فيها._\n"]
    buttons = []
    for i, acc in enumerate(accounts):
        info = await AccountManager.get_account_info(acc)
        label = f"{info['name']} — {info['phone']}"
        marker = " ✅ (الحالي)" if acc == current_poster else ""
        lines.append(f"{i + 1}. {label}{marker}")
        buttons.append([Button.inline(f"{'✅ ' if acc == current_poster else ''}📤 {label}", f"set_poster_do_{i}".encode())])

    buttons.append([Button.inline("↩️ رجوع", b"list_acc")])
    await event.edit("\n".join(lines), buttons=buttons, parse_mode="md")


@bot.on(events.CallbackQuery(data=re.compile(b"set_poster_do_(\\d+)")))
@owner_only
async def set_poster_do_handler(event):
    await event.answer()
    idx = int(event.data.decode().split("set_poster_do_")[1])
    accounts = db.get("accounts", [])
    if idx >= len(accounts):
        await event.edit("⚠️ الحساب غير موجود.", buttons=[nav_row()])
        return

    session_path = accounts[idx]
    info = await AccountManager.get_account_info(session_path)
    db["poster_session"] = session_path
    save_db(db)

    await event.edit(
        f"✅ **تم تعيين حساب الناشر:**\n"
        f"{info['name']} — {info['phone']}\n\n"
        f"_هذا الحساب سيُستخدم وحده لإرسال الروابط إلى قنوات الأرشيف._",
        buttons=[[Button.inline("👤 قائمة الحسابات", b"list_acc")], nav_row()],
        parse_mode="md",
    )


# ─────────────────────────────────────────────────────────────────────────────
# Step 2 — Create archive channels
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"make_ch"))
@owner_only
async def make_ch_handler(event):
    await event.answer("⏳ جاري إنشاء القنوات...")

    if not db["accounts"]:
        await event.edit(
            "🔒 **الخطوة 2 مقفلة**\n\nيجب ربط حساب أولاً (الخطوة 1) قبل إنشاء القنوات.",
            buttons=[[Button.inline("➕ ربط حساب ◄", b"add_acc")], nav_row()],
            parse_mode="md",
        )
        return

    existing = len(db.get("channels", {}))
    if existing >= 7:
        await event.edit(
            f"✅ **القنوات السبع موجودة بالفعل!**\n\n"
            + "\n".join(f"• {v}" for v in CHANNEL_KEYS.values())
            + "\n\n💡 إذا لم تستطع الحسابات الوصول للقنوات، استخدم روابط الدعوة للانضمام مباشرة.",
            buttons=[
                [Button.inline("🔗 انضمام عبر روابط دعوة", b"join_via_invites")],
                [Button.inline("➕ أضفني للقنوات كمسؤول", b"add_owner_to_ch")],
                [Button.inline("🔄 إعادة إنشاء القنوات", b"recreate_channels_confirm")],
                [Button.inline("🔁 إعادة الفرز من البداية", b"resort_from_scratch_confirm")],
                [Button.inline("⏭️ إضافة مصادر ◄", b"add_src")],
                nav_row(),
            ],
            parse_mode="md",
        )
        return

    auth_session = await _first_authorized_session(db["accounts"])
    if auth_session is None:
        await event.edit(
            "❌ **لا يوجد حساب مصرح به**\n\n"
            "جميع الجلسات منتهية الصلاحية. أعد ربط حساب واحد على الأقل من قائمة الحسابات.",
            buttons=[nav_row(b"add_acc")],
            parse_mode="md",
        )
        return

    await event.edit(
        "**②  إنشاء قنوات الأرشيف السبع**\n"
        "━━━━━━━━━━━━━━━━━━━━━\n"
        "⏳ جاري الإنشاء، يرجى الانتظار...\n\n"
        + "\n".join(f"🔄 {v}" for v in CHANNEL_KEYS.values()),
        parse_mode="md",
    )

    created = await create_archive_channels(auth_session, db, save_db)

    lines = ["✅ **تم إنشاء قنوات الأرشيف السبع:**\n"]
    for key, ch_id in created.items():
        title  = CHANNEL_KEYS.get(key, key)
        status = f"✅ `{ch_id}`" if isinstance(ch_id, int) else f"⚠️ {ch_id}"
        lines.append(f"• {title}: {status}")

    await event.edit(
        "\n".join(lines),
        buttons=[[Button.inline("⏭️ إضافة مصادر ◄", b"add_src")], nav_row(b"add_acc")],
        parse_mode="md",
    )


@bot.on(events.CallbackQuery(data=b"add_owner_to_ch"))
@owner_only
async def add_owner_to_ch_handler(event):
    await event.answer("⏳ جاري إضافتك للقنوات...")

    if not db.get("accounts"):
        await event.edit(
            "❌ لا يوجد حساب مرتبط. أضف حساباً أولاً.",
            buttons=[nav_row()],
        )
        return

    if not db.get("channels"):
        await event.edit(
            "❌ لا توجد قنوات مُنشأة بعد.",
            buttons=[nav_row()],
        )
        return

    await event.edit("⏳ جاري إضافتك كمسؤول في القنوات السبع، يرجى الانتظار...")

    try:
        await add_owner_to_channels(db)
        await event.edit(
            "✅ **تمت إضافتك كمسؤول في جميع القنوات السبع!**\n\n"
            "ستجد القنوات الآن في قائمة محادثاتك.",
            buttons=[[Button.inline("⏭️ إضافة مصادر ◄", b"add_src")], nav_row()],
            parse_mode="md",
        )
    except Exception as e:
        await event.edit(
            f"⚠️ حدث خطأ: {e}",
            buttons=[nav_row()],
        )


# ─────────────────────────────────────────────────────────────────────────────
# Join via invite links
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"join_via_invites"))
@owner_only
async def join_via_invites_handler(event):
    await event.answer()

    if not db.get("accounts"):
        await event.edit(
            "❌ لا يوجد حساب مرتبط. أضف حساباً أولاً.",
            buttons=[nav_row()],
        )
        return

    if not db.get("channels"):
        await event.edit(
            "❌ لا توجد قنوات مُنشأة بعد.",
            buttons=[nav_row()],
        )
        return

    stored_invites = db.get("channels_invites", [])
    msg_id  = event.message_id
    chat_id = OWNER_ID

    async def _edit(text: str, buttons=None):
        try:
            await bot.edit_message(chat_id, msg_id, text, buttons=buttons or [nav_row()], parse_mode="md")
        except Exception:
            pass

    if stored_invites:
        await _edit(
            f"🔗 **روابط الدعوة المحفوظة ({len(stored_invites)}):**\n"
            + "\n".join(f"• `{l}`" for l in stored_invites)
            + "\n\nأرسل **نعم** للاستخدام، أو أرسل الروابط الجديدة (كل رابط في سطر):",
        )
    else:
        await _edit(
            "🔗 **انضمام عبر روابط الدعوة**\n"
            "━━━━━━━━━━━━━━━━━━━━━\n\n"
            "أرسل روابط دعوة قنوات الأرشيف السبع (كل رابط في سطر):\n"
            "`https://t.me/+XXXXXXXXXXXX`",
        )

    try:
        async with bot.conversation(OWNER_ID, timeout=180, exclusive=False) as conv:
            ev    = await conv.wait_event(events.NewMessage(incoming=True, from_users=OWNER_ID))
            reply = ev.message
            text  = reply.text.strip()
            try:
                await reply.delete()
            except Exception:
                pass

            if text.lower() in ("نعم", "yes", "y") and stored_invites:
                invite_links = stored_invites
            else:
                invite_links = [l.strip() for l in text.split("\n") if "t.me/" in l]
                if not invite_links:
                    await _edit("❌ لم يتم العثور على روابط صالحة. حاول مرة أخرى.",
                                buttons=[[Button.inline("🔄 حاول مرة أخرى", b"join_via_invites")], nav_row()])
                    return
                db["channels_invites"] = invite_links
                save_db(db)

            accounts = db.get("accounts", [])
            await _edit(
                f"⏳ **جاري الانضمام عبر {len(invite_links)} رابط دعوة...**\n"
                f"الحسابات: {len(accounts)}\n\n"
                "قد يستغرق هذا بضع دقائق...",
            )

            summary = await join_accounts_via_invites(accounts, invite_links, db, save_db)

            lines = ["✅ **اكتمل الانضمام عبر روابط الدعوة:**\n"]
            for sess, s in summary.items():
                acc_name = sess.split("/")[-1]
                lines.append(
                    f"• حساب `{acc_name}`: "
                    f"✅ {s.get('joined',0)} انضم | "
                    f"♻️ {s.get('already',0)} موجود | "
                    f"❌ {s.get('errors',0)} خطأ"
                )
            hashes_found = len(db.get("channels_hashes", {}))
            lines.append(f"\n📌 هاش الوصول محفوظ لـ **{hashes_found}/7** قناة")

            await _edit(
                "\n".join(lines),
                buttons=[[Button.inline("⏭️ إضافة مصادر ◄", b"add_src")], nav_row()],
            )
    except asyncio.TimeoutError:
        await _edit("⏰ انتهت مهلة الإدخال. اضغط **انضمام عبر روابط دعوة** للمحاولة مجدداً.",
                    buttons=[[Button.inline("🔄 حاول مرة أخرى", b"join_via_invites")], nav_row()])
    except Exception:
        pass


# ─────────────────────────────────────────────────────────────────────────────
# Recreate channels — confirmation + execution
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"recreate_channels_confirm"))
@admin_only
async def recreate_channels_confirm_handler(event):
    await event.answer()
    await event.edit(
        "⚠️ **تحذير: إعادة إنشاء القنوات**\n\n"
        "سيتم:\n"
        "• حذف معرفات القنوات الحالية من الذاكرة\n"
        "• إنشاء 7 قنوات أرشيف جديدة\n\n"
        "⚠️ القنوات القديمة على تيليجرام **لن تُحذف** — فقط تُفقد الصلة بها.\n\n"
        "هل أنت متأكد؟",
        buttons=[
            [Button.inline("✅ نعم، أعد الإنشاء", b"recreate_channels_do")],
            [Button.inline("❌ إلغاء", b"make_ch")],
            nav_row(),
        ],
        parse_mode="md",
    )


@bot.on(events.CallbackQuery(data=b"recreate_channels_do"))
@admin_only
async def recreate_channels_do_handler(event):
    await event.answer("⏳ جاري إعادة الإنشاء...")

    if not db.get("accounts"):
        await event.edit("❌ لا يوجد حساب مرتبط. أضف حساباً أولاً.", buttons=[nav_row()])
        return

    auth_session = await _first_authorized_session(db["accounts"])
    if auth_session is None:
        await event.edit(
            "❌ **لا يوجد حساب مصرح به**\n\n"
            "جميع الجلسات منتهية الصلاحية. أعد ربط حساب واحد على الأقل قبل إعادة إنشاء القنوات.",
            buttons=[nav_row(b"add_acc")],
            parse_mode="md",
        )
        return

    # Clear old channel data
    db["channels"] = {}
    db.pop("channels_hashes", None)
    save_db(db)

    await event.edit(
        "🔄 **جاري إنشاء قنوات الأرشيف السبع من جديد...**\n"
        + "\n".join(f"🔄 {v}" for v in CHANNEL_KEYS.values()),
        parse_mode="md",
    )

    created = await create_archive_channels(auth_session, db, save_db)

    lines = ["✅ **تم إعادة إنشاء قنوات الأرشيف:**\n"]
    for key, ch_id in created.items():
        title  = CHANNEL_KEYS.get(key, key)
        status = f"✅ `{ch_id}`" if isinstance(ch_id, int) else f"⚠️ {ch_id}"
        lines.append(f"• {title}: {status}")

    lines.append(
        "\n💡 **التالي:** استخدم **🔗 انضمام عبر روابط دعوة** لإضافة باقي الحسابات."
    )

    await event.edit(
        "\n".join(lines),
        buttons=[
            [Button.inline("🔗 انضمام عبر روابط دعوة", b"join_via_invites")],
            [Button.inline("⏭️ إضافة مصادر ◄", b"add_src")],
            nav_row(),
        ],
        parse_mode="md",
    )


# ─────────────────────────────────────────────────────────────────────────────
# Re-sort from scratch (clears seen-set + progress, re-processes all links)
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"resort_from_scratch_confirm"))
@owner_only
async def resort_from_scratch_confirm_handler(event):
    await event.answer()
    raw_count  = get_raw_count()
    seen_count = get_seen_count()
    await event.edit(
        "⚠️ **إعادة الفرز من البداية**\n\n"
        "سيتم:\n"
        f"• مسح سجل الروابط المرئية ({seen_count:,} رابط)\n"
        f"• إعادة ضبط مؤشر التقدم (من 0 / {raw_count:,})\n"
        "• **حذف جميع الرسائل من قنوات الأرشيف**\n"
        "• إعادة فرز جميع الروابط وإرسالها للقنوات\n\n"
        "⚠️ هذا يفيد إذا كان الفرز السابق لم يُرسل الروابط للقنوات بسبب مشكلة في الوصول.\n\n"
        "هل أنت متأكد؟",
        buttons=[
            [Button.inline("✅ موافق", b"resort_from_scratch_do")],
            [Button.inline("❌ إلغاء", b"make_ch")],
            nav_row(),
        ],
        parse_mode="md",
    )


@bot.on(events.CallbackQuery(data=b"resort_from_scratch_do"))
@owner_only
async def resort_from_scratch_do_handler(event):
    await event.answer("⏳ جاري المسح وإعادة الضبط...")
    pmsg_id = event.message_id

    # 1) Clear the seen-set file and reset all stats (including per-channel)
    _do_clear_memory()

    # 2) Delete all messages from archive channels if accounts are linked
    if db.get("accounts") and db.get("channels"):
        await event.edit(
            "🗑 **جاري حذف الرسائل من قنوات الأرشيف...**\n"
            "━━━━━━━━━━━━━━━━━━━━━\n"
            "⏳ جاري التحضير...",
            parse_mode="md",
        )
        edit_cb = make_edit_callback(pmsg_id, OWNER_ID)
        results = await clear_archive_channels(db["accounts"], db, status_callback=edit_cb)
        total_deleted = sum(results.values())
    else:
        total_deleted = 0

    raw_count = get_raw_count()
    try:
        await bot.edit_message(
            OWNER_ID, pmsg_id,
            f"✅ **تم إعادة الضبط الكامل.**\n\n"
            f"• سجل الروابط المرئية: ممسوح\n"
            f"• مؤشر التقدم: 0 / {raw_count:,}\n"
            f"• رسائل الأرشيف المحذوفة: {total_deleted:,}\n\n"
            f"اضغط **⚡ فرز** لبدء الفرز من جديد وإرسال الروابط للقنوات.",
            buttons=[
                [Button.inline("⚡ بدء الفرز الآن ◄", b"run_sort")],
                nav_row(),
            ],
            parse_mode="md",
        )
    except Exception:
        pass


# ─────────────────────────────────────────────────────────────────────────────
# Step 3 — Add sources
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"add_src"))
@owner_only
async def add_src_handler(event):
    await event.answer()

    current = len(db.get("sources", []))
    msg_id  = event.message_id
    chat_id = OWNER_ID

    async def _edit(text: str, buttons=None):
        try:
            await bot.edit_message(chat_id, msg_id, text, buttons=buttons or [nav_row(b"make_ch")], parse_mode="md")
        except Exception:
            pass

    await _edit(
        f"**③  إضافة مصادر الروابط**\n"
        f"━━━━━━━━━━━━━━━━━━━━━\n"
        f"المصادر الحالية: **{current}**\n\n"
        f"📋 أرسل روابط مجموعات تيليجرام (كل رابط في سطر):\n"
        f"`https://t.me/medical_links_group`\n"
        f"`https://t.me/+AbCdEfGh1234`\n\n"
        f"📁 أو أرسل ملف `.txt` أو `.docx` يحتوي الروابط مباشرةً هنا.",
    )

    try:
        async with bot.conversation(event.sender_id, timeout=180, exclusive=False) as conv:
            _sender_id = event.sender_id
            _ev        = await conv.wait_event(events.NewMessage(incoming=True, from_users=_sender_id))
            links_msg  = _ev.message

            # ── Handle file upload (do NOT delete file messages) ───────────
            new_sources: list[str] = []
            if links_msg.document:
                fname = ""
                for attr in (links_msg.document.attributes or []):
                    if hasattr(attr, "file_name"):
                        fname = attr.file_name or ""
                        break
                ext = fname.lower().rsplit(".", 1)[-1] if "." in fname else ""
                if ext in ("txt", "docx"):
                    import tempfile as _tf, os as _os2, shutil as _sh2
                    tmp_dir  = _tf.mkdtemp()
                    tmp_path = _os2.path.join(tmp_dir, fname or f"sources.{ext}")
                    await _edit("⏳ جاري معالجة الملف...")
                    try:
                        await bot.download_media(links_msg, file=tmp_path)
                        if ext == "txt":
                            with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                                new_sources = [
                                    l.strip() for l in f.read().splitlines()
                                    if l.strip().startswith("http") or "t.me" in l
                                ]
                        else:
                            new_sources = _extract_links_from_docx(tmp_path)
                    finally:
                        _sh2.rmtree(tmp_dir, ignore_errors=True)
                else:
                    await _edit(
                        "❌ نوع الملف غير مدعوم. أرسل ملف `.txt` أو `.docx` فقط.",
                        buttons=[[Button.inline("🔄 حاول مرة أخرى", b"add_src")], nav_row()],
                    )
                    return
            elif links_msg.text and links_msg.text.strip():
                # Delete text messages (not file messages)
                try:
                    await links_msg.delete()
                except Exception:
                    pass
                new_sources = [l.strip() for l in links_msg.text.strip().split("\n") if l.strip()]

            if not new_sources:
                await _edit(
                    "⚠️ لم يُعثر على أي روابط. تأكد أن الروابط تبدأ بـ `https://t.me/`.",
                    buttons=[[Button.inline("🔄 حاول مرة أخرى", b"add_src")], nav_row()],
                )
                return

            added = 0
            for src in new_sources:
                if src not in db["sources"]:
                    db["sources"].append(src)
                    added += 1
            save_db(db)

            await _edit(
                f"✅ **تمت إضافة {added} مصدر جديد.**\n"
                f"📋 إجمالي المصادر: **{len(db['sources'])}**",
                buttons=[
                    [Button.inline("✏️ إضافة المزيد", b"add_src")],
                    [Button.inline("⏭️ بدء الحصاد ◄",  b"harvest")],
                    nav_row(b"make_ch"),
                ],
            )
    except asyncio.TimeoutError:
        await _edit("⏰ انتهت مهلة الإدخال.",
                    buttons=[[Button.inline("🔄 حاول مرة أخرى", b"add_src")], nav_row()])
    except Exception:
        pass


# ─────────────────────────────────────────────────────────────────────────────
# List sources
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"list_src"))
@owner_only
async def list_src_handler(event):
    await event.answer()
    if not db["sources"]:
        await event.edit(
            "❌ لا توجد مصادر مضافة بعد.",
            buttons=[[Button.inline("✏️ إضافة مصدر ◄", b"add_src")], nav_row()],
        )
        return

    lines = [f"📋 **المصادر ({len(db['sources'])}):**\n"]
    for i, src in enumerate(db["sources"], 1):
        lines.append(f"{i}. `{src}`")

    await event.edit(
        "\n".join(lines),
        buttons=[[Button.inline("✏️ إضافة مصادر جديدة", b"add_src")], nav_row()],
        parse_mode="md",
    )


# ─────────────────────────────────────────────────────────────────────────────
# Step 4 — Harvest
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"harvest"))
@owner_only
async def harvest_handler(event):
    await event.answer()

    if not db["accounts"]:
        await event.edit(
            "🔒 يجب ربط حساب أولاً.",
            buttons=[[Button.inline("➕ ربط حساب ◄", b"add_acc")], nav_row()],
        )
        return
    if not db["sources"]:
        await event.edit(
            "🔒 يجب إضافة مصادر أولاً.",
            buttons=[[Button.inline("✏️ إضافة مصدر ◄", b"add_src")], nav_row()],
        )
        return

    if sorter_ctrl.is_harvesting:
        await event.edit(
            "⚠️ **الحصاد يعمل بالفعل.**\n\nانتظر حتى ينتهي الحصاد الحالي، أو اضغط إيقاف.",
            buttons=[
                [Button.inline("⏹ إيقاف الحصاد", b"stop_harvest")],
                nav_row(),
            ],
            parse_mode="md",
        )
        return

    existing = get_raw_count()
    all_sessions = await _get_authorized_sessions(db["accounts"])
    if not all_sessions:
        await event.edit(
            "❌ **لا توجد حسابات متصلة.**\n\nجميع الحسابات انتهت جلساتها. أعد ربط حساب واحد على الأقل.",
            buttons=[[Button.inline("➕ ربط حساب", b"add_acc")], nav_row()],
            parse_mode="md",
        )
        return

    _stop_btn = [[Button.inline("⏹ إيقاف الحصاد", b"stop_harvest")]]
    prog_msg_id = event.message_id
    await event.edit(
        f"🌾 **الحصاد الموزع — جارٍ...**\n"
        f"━━━━━━━━━━━━━━━━━━━━━\n"
        f"👤 حسابات متصلة: **{len(all_sessions)}** | 📋 مصادر: **{len(db['sources'])}**\n"
        f"📦 روابط موجودة: **{existing:,}**\n\n"
        f"⏳ جاري الانضمام للمصادر وتوزيع العمل...",
        buttons=_stop_btn,
        parse_mode="md",
    )
    prog_cb = make_edit_callback(prog_msg_id, OWNER_ID, fixed_buttons=_stop_btn)

    sorter_ctrl.start_harvest()
    try:
        harvested = await harvest_sources(
            status_callback=prog_cb,
            db=db,
            sessions=all_sessions,
        )
    finally:
        sorter_ctrl.end_harvest()

    new_count = len(harvested) - existing
    try:
        await bot.edit_message(
            OWNER_ID, prog_msg_id,
            f"🎉 **اكتمل الحصاد!**\n"
            f"━━━━━━━━━━━━━━━━━━━━━\n"
            f"📦 إجمالي الروابط: **{len(harvested):,}**\n"
            f"🆕 روابط جديدة: **{new_count:,}**\n\n"
            f"💾 محفوظة في `raw_links.json`\nهل تريد بدء الفرز الآن؟",
            buttons=[
                [Button.inline("⏭️ بدء الفرز الآن ◄", b"run_sort")],
                nav_row(b"add_src"),
            ],
            parse_mode="md",
        )
    except Exception:
        pass


@bot.on(events.CallbackQuery(data=b"stop_harvest"))
@owner_only
async def stop_harvest_handler(event):
    await event.answer("⏹ جاري إيقاف الحصاد بعد الرسالة الحالية...")
    sorter_ctrl.stop_harvest()
    await event.answer("⏹ تم إرسال إشارة الإيقاف — سيتوقف الحصاد وتُحفظ الروابط.", alert=True)


# ─────────────────────────────────────────────────────────────────────────────
# Step 5 — Sort
# ─────────────────────────────────────────────────────────────────────────────

def _sort_control_buttons(paused: bool = False) -> list:
    """Inline buttons shown while sorting is active."""
    if paused:
        return [
            [Button.inline("▶️ استئناف", b"sort_resume"),
             Button.inline("⏹ إيقاف وحفظ", b"sort_stop")],
        ]
    return [
        [Button.inline("⏸ إيقاف مؤقت", b"sort_pause"),
         Button.inline("⏹ إيقاف وحفظ", b"sort_stop")],
    ]


@bot.on(events.CallbackQuery(data=b"run_sort"))
@owner_only
async def run_sort_handler(event):
    await event.answer()

    if not db["accounts"]:
        await event.edit(
            "🔒 يجب ربط حساب أولاً.",
            buttons=[[Button.inline("➕ ربط حساب ◄", b"add_acc")], nav_row()],
        )
        return
    if not db.get("channels"):
        await event.edit(
            "🔒 يجب إنشاء القنوات أولاً.",
            buttons=[[Button.inline("📺 إنشاء القنوات ◄", b"make_ch")], nav_row()],
        )
        return

    if sorter_ctrl.is_harvesting:
        await event.edit(
            "⚠️ **الحصاد يعمل الآن.** لا يمكن تشغيل الفرز في نفس الوقت.\n\nانتظر حتى ينتهي الحصاد.",
            buttons=[nav_row()],
            parse_mode="md",
        )
        return

    raw = load_raw_links()
    if not raw:
        await event.edit(
            "🔒 لا توجد روابط. قم بتشغيل الحصاد أولاً.",
            buttons=[[Button.inline("🌾 بدء الحصاد ◄", b"harvest")], nav_row()],
        )
        return

    sorter_ctrl.reset()

    start_from = db.get("progress", {}).get("last_sorted_index", 0)
    pct   = int(start_from / len(raw) * 100) if raw else 0
    bar_f = "▓" * (pct // 10)
    bar_e = "░" * (10 - pct // 10)

    # Edit the SAME message into the persistent progress message
    prog_msg_id = event.message_id
    await event.edit(
        f"📊 **الفرز الشامل — جارٍ...**\n"
        f"[{bar_f}{bar_e}] {pct}%\n\n"
        f"تم: **{start_from:,}** / {len(raw):,} رابط\n"
        f"{'🔄 استئناف من حيث توقفنا...' if start_from > 0 else '⚡ بدء الفرز الشامل...'}\n\n"
        f"_الأرقام تُحدَّث تلقائياً._",
        buttons=_sort_control_buttons(),
        parse_mode="md",
    )
    sorter_ctrl.set_progress_msg(prog_msg_id, OWNER_ID)

    authorized_sessions = await _get_authorized_sessions(db["accounts"])
    if not authorized_sessions:
        try:
            await bot.edit_message(
                OWNER_ID, prog_msg_id,
                "❌ **لا توجد حسابات متصلة.** أعد ربط حساب واحد على الأقل.",
                buttons=[nav_row(b"add_acc")],
                parse_mode="md",
            )
        except Exception:
            pass
        return

    await run_sorter(
        status_callback=status_msg,
        db=db,
        accounts=authorized_sessions,
        bot_client=bot,
        start_from=start_from,
    )
    sorter_ctrl.clear_progress_msg()

    if not sorter_ctrl.is_stopped():
        s = db["stats"]
        try:
            await bot.edit_message(
                OWNER_ID, prog_msg_id,
                f"🎯 **اكتمل الفرز بنجاح!**\n\n"
                f"📢 قنوات: **{s.get('ch_channels', 0):,}**  "
                f"👥 مجموعات: **{s.get('ch_groups', 0):,}**  "
                f"🤖 بوتات: **{s.get('ch_bots', 0):,}**\n"
                f"🔐 دعوات: **{s.get('ch_invite', 0):,}**  "
                f"📂 مجلدات: **{s.get('ch_addlist', 0):,}**  "
                f"🌐 غير طبي: **{s.get('ch_other', 0):,}**\n"
                f"💀 تالفة: **{s.get('ch_broken', 0):,}**",
                buttons=[
                    [Button.inline("🧠 اكتشاف ذكي ◄", b"smart_discover"),
                     Button.inline("🤝 انضمام ذكي ◄",  b"smart_join")],
                    nav_row(),
                ],
                parse_mode="md",
            )
        except Exception:
            pass


@bot.on(events.CallbackQuery(data=b"sort_pause"))
@owner_only
async def sort_pause_handler(event):
    await event.answer("⏸ سيتوقف بعد الدفعة الحالية...")
    sorter_ctrl.pause()
    # Update the progress message buttons to show Resume instead of Pause
    if sorter_ctrl.progress_msg_id and sorter_ctrl.progress_chat_id:
        try:
            await bot.edit_message(
                sorter_ctrl.progress_chat_id,
                sorter_ctrl.progress_msg_id,
                buttons=_sort_control_buttons(paused=True),
            )
        except Exception:
            pass
    await event.answer("⏸ الفرز سيتوقف بعد الدفعة الحالية.", alert=True)


@bot.on(events.CallbackQuery(data=b"sort_resume"))
@owner_only
async def sort_resume_handler(event):
    await event.answer("▶️ جاري الاستئناف...")
    sorter_ctrl.resume()
    if sorter_ctrl.progress_msg_id and sorter_ctrl.progress_chat_id:
        try:
            await bot.edit_message(
                sorter_ctrl.progress_chat_id,
                sorter_ctrl.progress_msg_id,
                buttons=_sort_control_buttons(paused=False),
            )
        except Exception:
            pass


@bot.on(events.CallbackQuery(data=b"sort_stop"))
@owner_only
async def sort_stop_handler(event):
    await event.answer("⏹ جاري الإيقاف وحفظ التقدم...")
    sorter_ctrl.stop()
    await event.answer("⏹ تم إيقاف الفرز. التقدم محفوظ — يمكن الاستئناف لاحقاً.", alert=True)


# ─────────────────────────────────────────────────────────────────────────────
# Step 6 — Smart Discovery
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"smart_discover"))
@owner_only
async def smart_discover_handler(event):
    await event.answer()

    if not db["accounts"]:
        await event.edit(
            "🔒 يجب ربط حساب أولاً.",
            buttons=[[Button.inline("➕ ربط حساب ◄", b"add_acc")], nav_row()],
        )
        return

    raw_count = get_raw_count()
    await event.edit(
        f"**⑥  الاكتشاف الذكي**\n"
        f"━━━━━━━━━━━━━━━━━━━━━\n"
        f"📦 الروابط الحالية: {raw_count:,}\n\n"
        f"🧠 يبحث بـ **8 طرق:**\n"
        f"  1️⃣ بحث بكلمات مفتاحية (100+ استعلام)\n"
        f"  2️⃣ قنوات مشابهة (Telegram AI)\n"
        f"  3️⃣ روابط من البيو والوصف\n"
        f"  4️⃣ روابط من الرسائل الأخيرة\n"
        f"  5️⃣ أنماط أسماء المستخدمين (GCC)\n"
        f"  6️⃣ مصفوفة الجهات × التخصصات 🆕\n"
        f"  7️⃣ بحث بالهاشتاقات الطبية 🆕\n"
        f"  8️⃣ Google Dorks (site:t.me) 🆕\n\n"
        f"هل تريد البدء؟",
        buttons=[
            [Button.inline("🚀 ابدأ الاكتشاف الآن", b"confirm_discover")],
            nav_row(b"run_sort"),
        ],
        parse_mode="md",
    )


@bot.on(events.CallbackQuery(data=b"confirm_discover"))
@owner_only
async def confirm_discover_handler(event):
    await event.answer("🧠 جاري الاكتشاف...")

    archive_ids  = {k: v for k, v in db.get("channels", {}).items() if isinstance(v, int)}
    source_links = db.get("sources", [])

    prog_msg_id = event.message_id
    await event.edit(
        "🧠 **الاكتشاف الذكي — جارٍ...**\n"
        "━━━━━━━━━━━━━━━━━━━━━\n"
        "⏳ جاري التحضير...",
        parse_mode="md",
    )
    prog_cb = make_edit_callback(prog_msg_id, OWNER_ID)

    authorized_disc = await _get_authorized_sessions(db["accounts"])
    if not authorized_disc:
        try:
            await bot.edit_message(OWNER_ID, prog_msg_id, "❌ لا توجد حسابات متصلة.",
                                   buttons=[nav_row()], parse_mode="md")
        except Exception:
            pass
        return

    new_count = await run_smart_discovery(
        status_callback=prog_cb,
        db=db,
        accounts=authorized_disc,
        archive_channel_ids=archive_ids,
        source_links=source_links,
    )

    result_text = (
        f"🎉 اكتُشف **{new_count:,}** رابط جديد!\nهل تريد الفرز الآن؟"
        if new_count > 0
        else "ℹ️ لم يُكتشف روابط جديدة. المصادر الحالية قد تكون مستنفدة."
    )
    result_btns = (
        [[Button.inline("⏭️ فرز الروابط الجديدة ◄", b"run_sort")], nav_row(b"run_sort")]
        if new_count > 0
        else [[Button.inline("⏭️ انضمام ذكي ◄", b"smart_join")], nav_row(b"run_sort")]
    )
    try:
        await bot.edit_message(OWNER_ID, prog_msg_id, result_text,
                               buttons=result_btns, parse_mode="md")
    except Exception:
        pass


# ─────────────────────────────────────────────────────────────────────────────
# Step 7 — Smart Join
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"smart_join"))
@owner_only
async def smart_join_handler(event):
    await event.answer()

    if not db["accounts"]:
        await event.edit(
            "🔒 يجب ربط حساب أولاً.",
            buttons=[[Button.inline("➕ ربط حساب ◄", b"add_acc")], nav_row()],
        )
        return
    if not db.get("channels"):
        await event.edit(
            "🔒 يجب إنشاء القنوات وتشغيل الفرز أولاً.",
            buttons=[[Button.inline("📺 إنشاء القنوات ◄", b"make_ch")], nav_row()],
        )
        return

    key_map = {
        "channels": (b"jch_channels", "📢 قناة القنوات"),
        "groups":   (b"jch_groups",   "👥 قناة المجموعات"),
        "invite":   (b"jch_invite",   "🔐 روابط الدعوة"),
        "addlist":  (b"jch_addlist",  "📂 المجلدات"),
        "bots":     (b"jch_bots",     "🤖 البوتات"),
        "other":    (b"jch_other",    "🌐 الروابط الأخرى"),
    }
    channel_buttons = []
    for key, (cb_data, label) in key_map.items():
        if key in db.get("channels", {}):
            channel_buttons.append([Button.inline(label, cb_data)])

    if not channel_buttons:
        await event.edit(
            "🔒 لا توجد قنوات أرشيف منشأة بعد.",
            buttons=[[Button.inline("📺 إنشاء القنوات ◄", b"make_ch")], nav_row()],
        )
        return

    joined_count = len(db.get("joined_links", []))
    await event.edit(
        f"**⑦  الانضمام الذكي**\n"
        f"━━━━━━━━━━━━━━━━━━━━━\n"
        f"🤝 تم الانضمام حتى الآن: **{joined_count:,}**\n\n"
        f"اختر القناة التي تريد الانضمام إلى روابطها:",
        buttons=channel_buttons + [nav_row(b"smart_discover")],
        parse_mode="md",
    )


async def _ask_join_count_and_start(event, source_key: str):
    await event.answer()
    channel_label = CHANNEL_KEYS.get(source_key, source_key)
    msg_id  = event.message_id
    chat_id = OWNER_ID

    async def _edit(text: str, buttons=None):
        try:
            await bot.edit_message(chat_id, msg_id, text,
                                   buttons=buttons or [nav_row(b"smart_join")], parse_mode="md")
        except Exception:
            pass

    await _edit(
        f"✅ المصدر: **{channel_label}**\n\n"
        f"📊 كم رابطاً تريد الانضمام إليه؟ (مثال: `20`)",
    )

    max_joins = None
    try:
        async with bot.conversation(event.sender_id, timeout=120, exclusive=False) as conv:
            _ev       = await conv.wait_event(events.NewMessage(incoming=True, from_users=event.sender_id))
            count_msg = _ev.message
            try:
                await count_msg.delete()
            except Exception:
                pass
            try:
                max_joins = int(count_msg.text.strip())
                if max_joins <= 0:
                    raise ValueError
            except ValueError:
                await _edit(
                    "❌ رقم غير صحيح. تم الإلغاء.",
                    buttons=[[Button.inline("🔄 حاول مرة أخرى", b"smart_join")], nav_row()],
                )
                return
    except asyncio.TimeoutError:
        await _edit("⏰ انتهت المهلة.", buttons=[[Button.inline("🔄 حاول مرة أخرى", b"smart_join")], nav_row()])
        return
    except Exception:
        return

    if max_joins is None:
        return

    await _edit(
        f"⏳ سيبدأ الانضمام إلى **{max_joins}** رابط من **{channel_label}**\n"
        f"🛡 نظام الحماية الذكي مفعّل (دفعات متقطعة)...",
    )

    source_ch_id  = db["channels"].get(source_key)
    links_to_join = []

    authorized_join = await _get_authorized_sessions(db["accounts"])
    if not authorized_join:
        await _edit("❌ لا توجد حسابات متصلة. أعد ربط حساب واحد على الأقل.")
        return

    if source_ch_id:
        try:
            async with TelegramClient(authorized_join[0], API_ID, API_HASH) as client:
                if not await client.is_user_authorized():
                    await _edit("❌ الجلسة الأولى غير مصرح بها. أعد ربط الحساب.")
                    return
                async for msg in client.iter_messages(int(source_ch_id), limit=500):
                    if msg.text:
                        found = re.findall(r"https?://t\.me/[\+a-zA-Z0-9_/]+", msg.text)
                        for lnk in found:
                            lnk = lnk.strip().rstrip("/")
                            if lnk not in links_to_join:
                                links_to_join.append(lnk)
                        if len(links_to_join) >= max_joins * 3:
                            break
        except EOFError:
            await _edit("❌ خطأ EOF في الجلسة — الجلسة تحتاج إعادة مصادقة.")
            return
        except Exception as e:
            await _edit(f"❌ خطأ في قراءة روابط القناة: {e}")

    if not links_to_join:
        await _edit(
            "❌ لا روابط في القناة.\nتأكد من تشغيل الفرز أولاً.",
            buttons=[[Button.inline("⚡ تشغيل الفرز أولاً ◄", b"run_sort")], nav_row()],
        )
        return

    await _edit(
        f"🤝 **الانضمام الذكي — جارٍ...**\n"
        f"━━━━━━━━━━━━━━━━━━━━━\n"
        f"[░░░░░░░░░░] 0%\n"
        f"📋 روابط للانضمام: **{min(max_joins, len(links_to_join))}**\n\n"
        f"⏳ جاري التحضير...",
    )
    join_cb = make_edit_callback(msg_id, OWNER_ID)

    await run_smart_joiner(
        status_callback=join_cb,
        links_to_join=links_to_join,
        accounts=authorized_join,
        db=db,
        max_joins=max_joins,
    )

    try:
        await bot.edit_message(OWNER_ID, msg_id, "🤝 **الانضمام الذكي — اكتمل ✅**",
                               buttons=[nav_row(b"smart_join")])
    except Exception:
        pass


@bot.on(events.CallbackQuery(data=b"jch_channels"))
@owner_only
async def jch_channels(event):
    await _ask_join_count_and_start(event, "channels")

@bot.on(events.CallbackQuery(data=b"jch_groups"))
@owner_only
async def jch_groups(event):
    await _ask_join_count_and_start(event, "groups")

@bot.on(events.CallbackQuery(data=b"jch_invite"))
@owner_only
async def jch_invite(event):
    await _ask_join_count_and_start(event, "invite")

@bot.on(events.CallbackQuery(data=b"jch_addlist"))
@owner_only
async def jch_addlist(event):
    await _ask_join_count_and_start(event, "addlist")

@bot.on(events.CallbackQuery(data=b"jch_bots"))
@owner_only
async def jch_bots(event):
    await _ask_join_count_and_start(event, "bots")

@bot.on(events.CallbackQuery(data=b"jch_other"))
@owner_only
async def jch_other(event):
    await _ask_join_count_and_start(event, "other")


# ─────────────────────────────────────────────────────────────────────────────
# Stats
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"stats"))
@owner_only
async def stats_handler(event):
    await event.answer()
    stats = db.get("stats", {})
    seen  = get_seen_count()
    raw   = get_raw_count()
    last  = db.get("progress", {}).get("last_sorted_index", 0)
    pct   = int(last / raw * 100) if raw else 0

    bar_filled = "█" * (pct // 10)
    bar_empty  = "░" * (10 - pct // 10)

    text = (
        "📊 **إحصائيات النظام**\n"
        "━━━━━━━━━━━━━━━━━━━━━\n\n"
        f"**الفرز:** {bar_filled}{bar_empty} {pct}%\n"
        f"  ↳ {last:,} / {raw:,} رابط\n\n"
        f"📦 **الروابط الخام:** {raw:,}\n"
        f"🔍 **المفحوصة:** {seen:,}\n"
        f"✅ **المرتبة بنجاح:** {stats.get('total_sorted', 0):,}\n"
        f"💀 **تالفة (منتهية فعلاً):** {stats.get('total_broken', 0):,}\n"
        f"🔐 **دعوات خاصة:** {stats.get('total_invite', 0):,}\n"
        f"⏭️ **المتخطاة (مكررة):** {stats.get('total_skipped_duplicate', 0):,}\n"
        f"🤝 **تم الانضمام إليها:** {len(db.get('joined_links', [])):,}\n\n"
        f"👤 **الحسابات:** {len(db.get('accounts', []))}\n"
        f"🔗 **المصادر:** {len(db.get('sources', []))}\n"
        f"📺 **القنوات:** {len(db.get('channels', {}))}/7\n\n"
        f"ℹ️ **ملاحظة عن التالفة:**\n"
        f"الروابط التالفة = يوزرنيم محذوف أو حساب غير موجود.\n"
        f"روابط +invite التي لم ينضم إليها الحساب → قناة 🔐 الدعوات."
    )

    await event.edit(text, buttons=[nav_row()], parse_mode="md")


# ─────────────────────────────────────────────────────────────────────────────
# Clear memory
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(data=b"clear_mem"))
@owner_only
async def clear_mem_handler(event):
    await event.answer()
    has_channels = bool(db.get("channels"))
    buttons = [
        [Button.inline("🧹 مسح الذاكرة فقط", b"confirm_clear")],
    ]
    if has_channels:
        buttons.append([Button.inline("🗑 مسح الذاكرة + حذف من القنوات", b"confirm_clear_full")])
    buttons.append([Button.inline("❌ إلغاء", b"home")])
    await event.edit(
        "⚠️ **تحذير — مسح الذاكرة**\n"
        "━━━━━━━━━━━━━━━━━━━━━\n\n"
        "سيتم مسح:\n"
        "• ذاكرة الروابط المفحوصة\n"
        "• مؤشر التقدم في الفرز\n"
        "• الإحصائيات\n\n"
        "الخيار الثاني يحذف **جميع الرسائل** من قنوات الأرشيف أيضاً.\n\n"
        "**هل أنت متأكد؟**",
        buttons=buttons,
        parse_mode="md",
    )


def _do_clear_memory():
    """Reset seen links, stats, and sort progress in-place."""
    clear_seen()
    db["stats"] = {
        "total_found": 0, "total_sorted": 0,
        "total_broken": 0, "total_skipped_duplicate": 0, "total_invite": 0,
        "ch_channels": 0, "ch_groups": 0, "ch_broken": 0,
        "ch_invite": 0, "ch_addlist": 0, "ch_bots": 0, "ch_other": 0,
    }
    db["progress"]["last_sorted_index"] = 0
    save_db(db)


@bot.on(events.CallbackQuery(data=b"confirm_clear"))
@owner_only
async def confirm_clear_handler(event):
    await event.answer()
    _do_clear_memory()
    await event.edit(
        "✅ **تم مسح الذاكرة.**\nيمكنك الآن بدء الفرز من جديد.",
        buttons=[[Button.inline("⚡ بدء الفرز ◄", b"run_sort")], nav_row()],
        parse_mode="md",
    )


@bot.on(events.CallbackQuery(data=b"confirm_clear_full"))
@owner_only
async def confirm_clear_full_handler(event):
    await event.answer()
    if not db.get("accounts"):
        _do_clear_memory()
        await event.edit(
            "✅ **تم مسح الذاكرة.**\n"
            "_(لا توجد حسابات مرتبطة — لم تُحذف رسائل القنوات)_",
            buttons=[nav_row()],
            parse_mode="md",
        )
        return

    _do_clear_memory()
    prog_msg_id = event.message_id

    await event.edit(
        "🗑 **جاري حذف الرسائل من قنوات الأرشيف...**\n"
        "━━━━━━━━━━━━━━━━━━━━━\n"
        "⏳ جاري التحضير...",
        parse_mode="md",
    )
    edit_cb = make_edit_callback(prog_msg_id, OWNER_ID)

    results = await clear_archive_channels(db["accounts"], db, status_callback=edit_cb)

    total_deleted = sum(results.values())
    summary_lines = [f"  • {k}: {v} رسالة" for k, v in results.items()]
    try:
        await bot.edit_message(
            OWNER_ID, prog_msg_id,
            f"✅ **تم مسح الذاكرة وحذف رسائل الأرشيف.**\n\n"
            f"🗑 إجمالي الرسائل المحذوفة: **{total_deleted:,}**\n"
            + "\n".join(summary_lines) + "\n\n"
            f"يمكنك الآن بدء الفرز من جديد.",
            buttons=[[Button.inline("⚡ بدء الفرز ◄", b"run_sort")], nav_row()],
            parse_mode="md",
        )
    except Exception:
        pass


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

# ─────────────────────────────────────────────────────────────────────────────
# Access Request System
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery(pattern=rb"req_(\d+)"))
async def request_access_handler(event):
    await event.answer()
    user_id = int(event.data.decode().split("_", 1)[1])

    if user_id != event.sender_id:
        await event.answer("❌ غير صالح.", alert=True)
        return

    db.update(load_db())
    if _is_authorized(user_id):
        await event.edit("✅ لديك وصول بالفعل. أرسل /start.", parse_mode="md")
        return

    sender = await event.get_sender()
    name     = getattr(sender, "first_name", "") or str(user_id)
    username = f"@{sender.username}" if getattr(sender, "username", None) else "بدون يوزرنيم"

    db.setdefault("pending_requests", {})[str(user_id)] = {
        "name": name,
        "username": username,
    }
    save_db(db)

    try:
        await bot.send_message(
            OWNER_ID,
            f"🔔 **طلب وصول جديد**\n\n"
            f"👤 الاسم: {name}\n"
            f"🆔 يوزرنيم: {username}\n"
            f"🔢 المعرف: `{user_id}`",
            buttons=[
                [Button.inline("✅ قبول", f"ap_{user_id}".encode()),
                 Button.inline("❌ رفض",  f"dn_{user_id}".encode())],
            ],
            parse_mode="md",
        )
    except Exception:
        pass

    await event.edit(
        "✅ **تم إرسال طلبك للمالك.**\n\nسيتم إشعارك فور البت في طلبك.",
        parse_mode="md",
    )


@bot.on(events.CallbackQuery(pattern=rb"ap_(\d+)"))
@admin_only
async def approve_user_handler(event):
    user_id = int(event.data.decode().split("_", 1)[1])
    db.update(load_db())

    if user_id not in db.setdefault("trusted_users", []):
        db["trusted_users"].append(user_id)

    info = db.setdefault("pending_requests", {}).pop(str(user_id), {})
    save_db(db)

    try:
        await bot.send_message(
            user_id,
            "✅ **تم قبول طلبك!**\n\nأرسل /start لفتح لوحة التحكم.",
            parse_mode="md",
        )
    except Exception:
        pass

    name = info.get("name", str(user_id))
    await event.answer(f"✅ تم قبول {name}")
    await event.edit(
        f"✅ **تم قبول الوصول**\n👤 {name} | `{user_id}`",
        parse_mode="md",
    )


@bot.on(events.CallbackQuery(pattern=rb"dn_(\d+)"))
@admin_only
async def deny_user_handler(event):
    user_id = int(event.data.decode().split("_", 1)[1])
    db.update(load_db())

    info = db.setdefault("pending_requests", {}).pop(str(user_id), {})
    save_db(db)

    try:
        await bot.send_message(
            user_id,
            "❌ **تم رفض طلبك.**\n\nللاستفسار تواصل مع المالك.",
            parse_mode="md",
        )
    except Exception:
        pass

    name = info.get("name", str(user_id))
    await event.answer(f"❌ تم رفض {name}")
    await event.edit(
        f"❌ **تم رفض الطلب**\n👤 {name} | `{user_id}`",
        parse_mode="md",
    )


@bot.on(events.NewMessage(pattern="/trusted"))
@admin_only
async def trusted_users_handler(event):
    db.update(load_db())
    trusted = db.get("trusted_users", [])

    if not trusted:
        await event.respond(
            "👥 **المستخدمون الموثوقون**\n\nلا يوجد مستخدمون موثوقون حتى الآن.\n\n"
            "ستصلك إشعارات عندما يطلب أحدهم الوصول.",
            parse_mode="md",
        )
        return

    lines = [f"👥 **المستخدمون الموثوقون ({len(trusted)}):**\n"]
    for uid in trusted:
        lines.append(f"• `{uid}`")

    lines.append("\n_لإزالة مستخدم: أرسل /revoke <ID>_")
    await event.respond("\n".join(lines), parse_mode="md")


@bot.on(events.NewMessage(pattern=r"/revoke (\d+)"))
@admin_only
async def revoke_user_handler(event):
    db.update(load_db())
    user_id = int(event.pattern_match.group(1))

    if user_id in db.get("trusted_users", []):
        db["trusted_users"].remove(user_id)
        save_db(db)
        await event.respond(f"✅ تم إلغاء وصول المستخدم `{user_id}`.", parse_mode="md")
        try:
            await bot.send_message(user_id, "🚫 تم إلغاء وصولك من قِبل المالك.", parse_mode="md")
        except Exception:
            pass
    else:
        await event.respond(f"❌ المعرف `{user_id}` غير موجود في قائمة الموثوقين.", parse_mode="md")


# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.NewMessage(pattern="/whoami"))
async def whoami_handler(event):
    """No auth check — lets you find your real Telegram ID."""
    if _is_stale_event(event):
        return
    await event.respond(
        f"🪪 **معرفك في تيليجرام:**\n`{event.sender_id}`\n\n"
        f"📌 **OWNER_ID المحمّل:** `{OWNER_ID}`\n\n"
        + ("✅ أنت المالك." if event.sender_id == OWNER_ID else
           "❌ لا تتطابق! غيّر OWNER_ID في الأسرار ليساوي معرّفك."),
        parse_mode="md",
    )


# ─────────────────────────────────────────────────────────────────────────────
# Catch-all fallback — MUST be the last registered handler.
# Answers any callback that no specific handler above matched,
# so buttons never get stuck in a permanent loading state after restarts.
# ─────────────────────────────────────────────────────────────────────────────

@bot.on(events.CallbackQuery())
async def fallback_callback_handler(event):
    try:
        await event.answer("🔄 أعد تشغيل البوت بـ /start", alert=False)
    except Exception:
        pass


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

async def _check_sessions_on_startup():
    """Check all saved accounts at startup and alert owner if any sessions are expired."""
    accounts = db.get("accounts", [])
    if not accounts:
        return
    expired = []
    for acc in accounts:
        info = await AccountManager.get_account_info(acc)
        if info.get("unauthorized"):
            expired.append(acc)
    if expired:
        try:
            await bot.send_message(
                OWNER_ID,
                f"🔴 **تنبيه: {len(expired)} حساب انتهت جلسته!**\n\n"
                f"السبب الأكثر شيوعاً: انتقل البوت إلى خادم جديد (IP مختلف) "
                f"فأوقف تيليجرام هذه الجلسات تلقائياً لحماية الحسابات.\n\n"
                f"**الحل:** افتح 👤 حساباتي ثم أعد ربط كل حساب منتهٍ.",
                buttons=[[Button.inline("👤 عرض الحسابات", b"list_acc")]],
                parse_mode="md",
            )
        except Exception:
            pass


async def main():
    global _BOT_START_TIME
    asyncio.create_task(_keep_alive_http())
    while True:
        try:
            # Record start time BEFORE connecting so any event whose timestamp
            # predates this moment is treated as a stale replay and silently dropped.
            _BOT_START_TIME = _time_module.time()
            await bot.start(bot_token=BOT_TOKEN)
            try:
                await bot(GetStateRequest())
            except Exception:
                pass
            print(f"🤖 البوت يعمل... OWNER_ID={OWNER_ID} | أرسل /start في تيليجرام للبدء.")
            asyncio.create_task(_check_sessions_on_startup())
            await bot.run_until_disconnected()
        except Exception as e:
            print(f"⚠️ انقطع الاتصال: {e} — إعادة المحاولة خلال 10 ثوانٍ...")
            await asyncio.sleep(10)
        finally:
            try:
                await bot.disconnect()
            except Exception:
                pass


# ─────────────────────────────────────────────────────────────────────────────
# File Import Handler — .txt / .docx link extraction  ★ NEW (Task 3)
# ─────────────────────────────────────────────────────────────────────────────

_FILE_LINK_RE = re.compile(
    r"(?:https?://)?t\.me/(?:joinchat/|\+)?[A-Za-z0-9_\-/]{5,}"
    r"|(?:https?://)?(?:chat\.whatsapp\.com|wa\.me)/[A-Za-z0-9_\-]+",
    re.IGNORECASE,
)

# Matches standalone @username handles (Telegram usernames are 5–32 chars,
# letters/digits/underscores, must start with a letter or digit).
# Negative lookbehind ensures we don't capture email addresses or mid-word @.
_AT_HANDLE_RE = re.compile(
    r"(?<![A-Za-z0-9_])@([A-Za-z][A-Za-z0-9_]{3,31})(?![A-Za-z0-9_@\.])",
)


def _extract_links_from_text(text: str) -> list[str]:
    """Extract all Telegram and WhatsApp links (including @username handles) from text."""
    found = []
    seen_norms: set[str] = set()

    # ── t.me URLs and WhatsApp links ────────────────────────────────────────
    for m in _FILE_LINK_RE.findall(text):
        link = m.strip().rstrip("/")
        if not link.startswith("http"):
            link = "https://" + link
        norm = link.lower()
        if norm not in seen_norms:
            found.append(link)
            seen_norms.add(norm)

    # ── @username handles → https://t.me/username ───────────────────────────
    for m in _AT_HANDLE_RE.finditer(text):
        username = m.group(1)
        link = f"https://t.me/{username}"
        norm = link.lower()
        if norm not in seen_norms:
            found.append(link)
            seen_norms.add(norm)

    return found


def _extract_links_from_docx(path: str) -> list[str]:
    """Extract all links from a .docx file (paragraphs + tables)."""
    try:
        import docx as _docx
    except ImportError:
        return []
    doc = _docx.Document(path)
    text_parts = []
    for para in doc.paragraphs:
        text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)
    return _extract_links_from_text("\n".join(text_parts))


@bot.on(events.NewMessage(pattern=None))
@owner_only
async def file_import_handler(event):
    """
    Handle uploaded .txt or .docx files.
    Extracts all Telegram/WhatsApp links, deduplicates against every known link,
    and appends new ones to raw_links.json for the normal sort pipeline.
    """
    msg = event.message
    if not msg.document:
        return

    # Check file extension
    fname = ""
    if msg.document.attributes:
        for attr in msg.document.attributes:
            if hasattr(attr, "file_name"):
                fname = attr.file_name or ""
                break

    ext = fname.lower().rsplit(".", 1)[-1] if "." in fname else ""
    if ext not in ("txt", "docx"):
        return  # ignore non-txt/docx documents silently

    status = await event.reply(
        f"📂 **استيراد الملف:** `{fname}`\n"
        "⏳ جارٍ تنزيل الملف وفحص الروابط..."
    )

    import tempfile, os as _os
    tmp_dir = tempfile.mkdtemp()
    tmp_path = _os.path.join(tmp_dir, fname or f"import.{ext}")

    try:
        await bot.download_media(msg, file=tmp_path)

        # ── Extract links ──────────────────────────────────────────────────
        if ext == "txt":
            with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                raw_text = f.read()
            found_links = _extract_links_from_text(raw_text)
        else:
            found_links = _extract_links_from_docx(tmp_path)

        if not found_links:
            await status.edit(
                f"📂 **{fname}**\n"
                "⚠️ لم يُعثر على أي روابط تيليجرام أو واتساب في الملف."
            )
            return

        # ── Deduplicate against ALL known links (raw + archived + joined) ──
        known = load_all_known_links(joined_links=db.get("joined_links", []))
        new_links: list[str] = []
        seen_this_file: set[str] = set()

        for link in found_links:
            norm = normalize_link(link)
            if norm not in known and norm not in seen_this_file:
                new_links.append(link)
                seen_this_file.add(norm)
                known.add(norm)

        # ── Append new links to raw pipeline ──────────────────────────────
        if new_links:
            existing = load_raw_links()
            save_raw_links(existing + new_links)

        dupes = len(found_links) - len(new_links)
        await status.edit(
            f"📂 **استيراد مكتمل: `{fname}`**\n\n"
            f"🔍 روابط مُكتشفة في الملف: **{len(found_links):,}**\n"
            f"✅ جديدة أُضيفت للمعالجة: **{len(new_links):,}**\n"
            f"♻️ تكرارات تجاهلها: **{dupes:,}**\n\n"
            + (
                "▶️ شغّل الفرز لأرشفة الروابط الجديدة."
                if new_links else
                "ℹ️ جميع الروابط موجودة مسبقاً — لا جديد."
            )
        )

    except Exception as e:
        await status.edit(f"❌ خطأ أثناء معالجة الملف:\n`{e}`")
    finally:
        try:
            import shutil as _sh
            _sh.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass


# ─────────────────────────────────────────────────────────────────────────────
# Inline text link sort — links pasted directly in chat
# ─────────────────────────────────────────────────────────────────────────────

_CHANNEL_KEY_EMOJI = {
    "channels": "📢",
    "groups":   "👥",
    "bots":     "🤖",
    "invite":   "🔐",
    "addlist":  "📂",
    "other":    "🌐",
    "broken":   "💀",
}


@bot.on(events.NewMessage(pattern=None))
@owner_only
async def inline_text_sort_handler(event):
    """
    Triggered when the owner sends a plain text message (no file) containing
    at least one t.me link.  Sorts each link immediately: deduplicates, posts
    to the correct archive channel, marks as seen, then deletes the original
    user message from chat.
    """
    msg = event.message

    # ── Ignore documents (handled by file_import_handler) ─────────────────
    if msg.document:
        return

    # ── Ignore bot commands ────────────────────────────────────────────────
    text = msg.text or msg.message or ""
    if text.startswith("/"):
        return

    # ── Extract links ──────────────────────────────────────────────────────
    links = _extract_links_from_text(text)
    if not links:
        return

    # ── Guard: need accounts & archive channels ────────────────────────────
    accounts = db.get("accounts", [])
    channels = db.get("channels", {})
    if not accounts:
        status = await event.reply(
            "⚠️ لا توجد حسابات مرتبطة — أضف حساباً أولاً قبل الفرز.",
            parse_mode="md",
        )
        return
    if len(channels) < 7:
        status = await event.reply(
            "⚠️ قنوات الأرشيف غير مهيأة — أنشئها من الخطوة 2 أولاً.",
            parse_mode="md",
        )
        return

    # ── Acknowledge immediately ────────────────────────────────────────────
    status = await event.reply(
        f"🔃 **جاري فرز {len(links)} رابط...**\n"
        "⏳ يُرجى الانتظار...",
        parse_mode="md",
    )
    user_msg_id = event.message.id

    # ── Sort ───────────────────────────────────────────────────────────────
    async def _upd(text_upd: str):
        try:
            await status.edit(text_upd, parse_mode="md")
        except Exception:
            pass

    result = await sort_links_inline(
        links    = links,
        db       = db,
        accounts = accounts,
        bot_client = bot,
        status_callback = _upd,
    )

    # ── Delete original user message ───────────────────────────────────────
    try:
        await bot.delete_messages(OWNER_ID, [user_msg_id])
    except Exception:
        pass

    # ── Build summary ──────────────────────────────────────────────────────
    posted   = result["posted"]
    known    = result["already_known"]
    errors   = result["errors"]
    total    = len(links)

    summary_lines = [f"✅ **اكتمل الفرز الفوري — {total} رابط**\n"]

    if posted:
        summary_lines.append(f"📬 **أُرشف ({len(posted)}):**")
        for lnk, ck in posted:
            emoji = _CHANNEL_KEY_EMOJI.get(ck, "📌")
            summary_lines.append(f"  {emoji} {lnk}")

    if known:
        summary_lines.append(f"\n♻️ **مُرشَف مسبقاً ({len(known)}) — تم تجاهله:**")
        for lnk in known:
            summary_lines.append(f"  • {lnk}")

    if errors:
        summary_lines.append(f"\n❌ **فشل الإرسال ({len(errors)}) — الرابط لم يصل للقناة:**")
        for item in errors:
            if isinstance(item, tuple):
                lnk, err = item
                summary_lines.append(f"  • {lnk}")
                if err:
                    summary_lines.append(f"    `{err[:120]}`")
            else:
                summary_lines.append(f"  • {item}")

    # If any links failed to send, offer the full sorter as a retry path
    extra_buttons = None
    if errors:
        extra_buttons = [[Button.inline("⚡ فرز شامل (إعادة المحاولة)", b"run_sort"),
                          Button.inline("🏠 القائمة", b"home")]]

    await status.edit("\n".join(summary_lines), buttons=extra_buttons, parse_mode="md")


if __name__ == "__main__":
    asyncio.run(main())
